"""
Document Parser for multiple file formats.

Supports:
- PDF documents (with and without text layer)
- DOCX documents
- HTML documents
- Plain text documents
- Images (PNG, JPG, TIFF)

Extracts text while preserving document structure.
"""

from typing import Dict, List, Any, Optional, Union
from dataclasses import dataclass, field
from pathlib import Path
import re

from loguru import logger

from .utils import (
    FileTypeDetector,
    TextCleaner,
    LanguageDetector,
    EncodingDetector,
    DocumentMetadataExtractor,
    TableExtractor,
)
from .ocr_engine import OCREngineFactory, OCRResult


@dataclass
class DocumentElement:
    """Represents a structured element in a document."""
    type: str  # heading, paragraph, list, table, figure, code, equation
    content: str
    level: Optional[int] = None  # For headings
    metadata: Dict[str, Any] = field(default_factory=dict)
    children: List['DocumentElement'] = field(default_factory=list)


@dataclass
class ParsedDocument:
    """Complete parsed document with structure and metadata."""
    text: str
    elements: List[DocumentElement]
    metadata: Dict[str, Any]
    structure: Dict[str, Any]
    raw_content: Optional[str] = None


class PDFParser:
    """Parse PDF documents with text extraction and OCR fallback."""

    def __init__(self, ocr_engine: Optional[Any] = None):
        """
        Initialize PDF parser.

        Args:
            ocr_engine: OCR engine for scanned PDFs
        """
        self.ocr_engine = ocr_engine

    def parse(self, file_path: Path) -> ParsedDocument:
        """
        Parse PDF document.

        Args:
            file_path: Path to PDF file

        Returns:
            Parsed document with structure
        """
        try:
            import pdfplumber
            use_pdfplumber = True
        except ImportError:
            import PyPDF2
            use_pdfplumber = False
            logger.info("pdfplumber not available, using PyPDF2")

        elements = []
        full_text = []
        metadata = {}

        try:
            if use_pdfplumber:
                parsed = self._parse_with_pdfplumber(file_path)
            else:
                parsed = self._parse_with_pypdf2(file_path)

            return parsed

        except Exception as e:
            logger.error(f"PDF parsing failed: {e}")
            # Fallback to OCR if text extraction fails
            if self.ocr_engine:
                logger.info("Falling back to OCR for PDF processing")
                return self._parse_with_ocr(file_path)
            raise

    def _parse_with_pdfplumber(self, file_path: Path) -> ParsedDocument:
        """Parse PDF using pdfplumber (better structure preservation)."""
        import pdfplumber

        elements = []
        full_text = []
        tables_found = []

        with pdfplumber.open(file_path) as pdf:
            # Extract metadata
            metadata = pdf.metadata or {}
            metadata.update({
                'num_pages': len(pdf.pages),
                'parser': 'pdfplumber',
            })

            # Process each page
            for page_num, page in enumerate(pdf.pages):
                # Extract text
                page_text = page.extract_text()

                if page_text:
                    # Add page text
                    full_text.append(page_text)

                    # Parse page structure
                    page_elements = self._parse_text_structure(
                        page_text,
                        page_num=page_num
                    )
                    elements.extend(page_elements)

                    # Extract tables
                    tables = page.extract_tables()
                    for table_idx, table in enumerate(tables):
                        if table:
                            table_md = TableExtractor.table_to_markdown(table)
                            elements.append(DocumentElement(
                                type='table',
                                content=table_md,
                                metadata={
                                    'page': page_num,
                                    'table_index': table_idx,
                                    'rows': len(table),
                                    'cols': len(table[0]) if table else 0,
                                }
                            ))
                            tables_found.append(table)

                else:
                    # No text extracted, might be scanned
                    logger.warning(f"No text on page {page_num}, may need OCR")

        combined_text = '\n\n'.join(full_text)

        return ParsedDocument(
            text=combined_text,
            elements=elements,
            metadata=metadata,
            structure={
                'total_elements': len(elements),
                'tables': len(tables_found),
                'parser': 'pdfplumber'
            },
            raw_content=combined_text
        )

    def _parse_with_pypdf2(self, file_path: Path) -> ParsedDocument:
        """Parse PDF using PyPDF2 (fallback)."""
        import PyPDF2

        elements = []
        full_text = []

        with open(file_path, 'rb') as file:
            reader = PyPDF2.PdfReader(file)

            # Extract metadata
            metadata = {
                'num_pages': len(reader.pages),
                'parser': 'pypdf2',
            }

            if reader.metadata:
                metadata.update({
                    'title': reader.metadata.get('/Title', ''),
                    'author': reader.metadata.get('/Author', ''),
                    'subject': reader.metadata.get('/Subject', ''),
                    'creator': reader.metadata.get('/Creator', ''),
                })

            # Extract text from each page
            for page_num, page in enumerate(reader.pages):
                page_text = page.extract_text()

                if page_text:
                    full_text.append(page_text)

                    # Parse structure
                    page_elements = self._parse_text_structure(
                        page_text,
                        page_num=page_num
                    )
                    elements.extend(page_elements)

        combined_text = '\n\n'.join(full_text)

        return ParsedDocument(
            text=combined_text,
            elements=elements,
            metadata=metadata,
            structure={
                'total_elements': len(elements),
                'parser': 'pypdf2'
            },
            raw_content=combined_text
        )

    def _parse_with_ocr(self, file_path: Path) -> ParsedDocument:
        """Parse scanned PDF using OCR."""
        try:
            from pdf2image import convert_from_path
        except ImportError:
            logger.error("pdf2image not installed. Install with: pip install pdf2image")
            raise

        # Convert PDF to images
        images = convert_from_path(file_path)

        elements = []
        full_text = []

        for page_num, image in enumerate(images):
            # Save temporary image
            import tempfile
            with tempfile.NamedTemporaryFile(suffix='.png', delete=False) as tmp:
                image.save(tmp.name)
                tmp_path = Path(tmp.name)

            # Run OCR
            ocr_result = self.ocr_engine.process_image(tmp_path)

            if ocr_result.text:
                full_text.append(ocr_result.text)

                # Parse structure
                page_elements = self._parse_text_structure(
                    ocr_result.text,
                    page_num=page_num
                )
                elements.extend(page_elements)

            # Cleanup
            tmp_path.unlink()

        combined_text = '\n\n'.join(full_text)

        return ParsedDocument(
            text=combined_text,
            elements=elements,
            metadata={
                'num_pages': len(images),
                'parser': 'ocr',
                'ocr_confidence': sum(e.metadata.get('confidence', 0) for e in elements) / len(elements) if elements else 0
            },
            structure={
                'total_elements': len(elements),
                'parser': 'ocr'
            },
            raw_content=combined_text
        )

    def _parse_text_structure(
        self,
        text: str,
        page_num: int = 0
    ) -> List[DocumentElement]:
        """
        Parse text into structured elements.

        Args:
            text: Plain text content
            page_num: Page number

        Returns:
            List of document elements
        """
        elements = []
        lines = text.split('\n')

        current_paragraph = []
        current_list = []
        in_list = False

        for line in lines:
            stripped = line.strip()

            if not stripped:
                # End current paragraph or list
                if current_paragraph:
                    elements.append(DocumentElement(
                        type='paragraph',
                        content='\n'.join(current_paragraph),
                        metadata={'page': page_num}
                    ))
                    current_paragraph = []

                if current_list:
                    elements.append(DocumentElement(
                        type='list',
                        content='\n'.join(current_list),
                        metadata={'page': page_num, 'list_type': 'unordered'}
                    ))
                    current_list = []
                    in_list = False

                continue

            # Check for heading (heuristic: short lines, possibly all caps)
            if len(stripped) < 100 and not stripped.endswith(('.', ',', ';', ':')):
                if stripped.isupper() or (stripped[0].isupper() and ':' not in stripped):
                    # Likely a heading
                    level = 1 if stripped.isupper() else 2
                    elements.append(DocumentElement(
                        type='heading',
                        content=stripped,
                        level=level,
                        metadata={'page': page_num}
                    ))
                    continue

            # Check for list item
            list_pattern = r'^[\-\*\•\◦]\s+|^\d+[\.\)]\s+'
            if re.match(list_pattern, stripped):
                in_list = True
                current_list.append(stripped)
                continue

            # Regular paragraph line
            if in_list:
                # End list
                elements.append(DocumentElement(
                    type='list',
                    content='\n'.join(current_list),
                    metadata={'page': page_num, 'list_type': 'unordered'}
                ))
                current_list = []
                in_list = False

            current_paragraph.append(line)

        # Add remaining content
        if current_paragraph:
            elements.append(DocumentElement(
                type='paragraph',
                content='\n'.join(current_paragraph),
                metadata={'page': page_num}
            ))

        if current_list:
            elements.append(DocumentElement(
                type='list',
                content='\n'.join(current_list),
                metadata={'page': page_num, 'list_type': 'unordered'}
            ))

        return elements


class DOCXParser:
    """Parse DOCX documents with full structure preservation."""

    def parse(self, file_path: Path) -> ParsedDocument:
        """
        Parse DOCX document.

        Args:
            file_path: Path to DOCX file

        Returns:
            Parsed document with structure
        """
        try:
            from docx import Document
        except ImportError:
            logger.error("python-docx not installed. Install with: pip install python-docx")
            raise

        doc = Document(file_path)

        elements = []
        full_text = []

        # Extract metadata
        metadata = {
            'parser': 'docx',
        }
        core_props = doc.core_properties
        if core_props:
            metadata.update({
                'title': core_props.title or '',
                'author': core_props.author or '',
                'subject': core_props.subject or '',
                'created': str(core_props.created) if core_props.created else '',
                'modified': str(core_props.modified) if core_props.modified else '',
            })

        # Process paragraphs
        for para in doc.paragraphs:
            text = para.text.strip()
            if not text:
                continue

            full_text.append(text)

            # Determine element type
            if para.style.name.startswith('Heading'):
                # Extract heading level
                level_match = re.search(r'Heading (\d+)', para.style.name)
                level = int(level_match.group(1)) if level_match else 1

                elements.append(DocumentElement(
                    type='heading',
                    content=text,
                    level=level,
                    metadata={'style': para.style.name}
                ))

            elif para.style.name.startswith('List'):
                elements.append(DocumentElement(
                    type='list',
                    content=text,
                    metadata={'style': para.style.name}
                ))

            else:
                # Regular paragraph
                elements.append(DocumentElement(
                    type='paragraph',
                    content=text,
                    metadata={'style': para.style.name}
                ))

        # Process tables
        for table_idx, table in enumerate(doc.tables):
            rows = []
            for row in table.rows:
                cells = [cell.text.strip() for cell in row.cells]
                rows.append(cells)

            if rows:
                table_md = TableExtractor.table_to_markdown(rows)
                elements.append(DocumentElement(
                    type='table',
                    content=table_md,
                    metadata={
                        'table_index': table_idx,
                        'rows': len(rows),
                        'cols': len(rows[0]) if rows else 0,
                    }
                ))

        combined_text = '\n\n'.join(full_text)

        return ParsedDocument(
            text=combined_text,
            elements=elements,
            metadata=metadata,
            structure={
                'total_elements': len(elements),
                'tables': sum(1 for e in elements if e.type == 'table'),
                'headings': sum(1 for e in elements if e.type == 'heading'),
                'parser': 'docx'
            },
            raw_content=combined_text
        )


class HTMLParser:
    """Parse HTML documents with structure preservation."""

    def parse(self, file_path: Path) -> ParsedDocument:
        """
        Parse HTML document.

        Args:
            file_path: Path to HTML file

        Returns:
            Parsed document with structure
        """
        try:
            from bs4 import BeautifulSoup
        except ImportError:
            logger.error("beautifulsoup4 not installed. Install with: pip install beautifulsoup4")
            raise

        # Read file
        content = EncodingDetector.read_text_file(file_path)

        soup = BeautifulSoup(content, 'html.parser')

        elements = []
        full_text = []

        # Extract metadata
        metadata = {'parser': 'html'}
        title_tag = soup.find('title')
        if title_tag:
            metadata['title'] = title_tag.text.strip()

        # Extract meta tags
        for meta in soup.find_all('meta'):
            name = meta.get('name', meta.get('property', ''))
            content_val = meta.get('content', '')
            if name and content_val:
                metadata[name] = content_val

        # Process headings
        for heading in soup.find_all(['h1', 'h2', 'h3', 'h4', 'h5', 'h6']):
            text = heading.get_text(strip=True)
            if text:
                level = int(heading.name[1])  # Extract number from h1, h2, etc.
                full_text.append(text)
                elements.append(DocumentElement(
                    type='heading',
                    content=text,
                    level=level,
                    metadata={'tag': heading.name}
                ))

        # Process paragraphs
        for para in soup.find_all('p'):
            text = para.get_text(strip=True)
            if text:
                full_text.append(text)
                elements.append(DocumentElement(
                    type='paragraph',
                    content=text,
                    metadata={'tag': 'p'}
                ))

        # Process lists
        for list_tag in soup.find_all(['ul', 'ol']):
            items = []
            for li in list_tag.find_all('li', recursive=False):
                items.append(li.get_text(strip=True))

            if items:
                list_text = '\n'.join(f"- {item}" for item in items)
                full_text.append(list_text)
                elements.append(DocumentElement(
                    type='list',
                    content=list_text,
                    metadata={
                        'tag': list_tag.name,
                        'list_type': 'ordered' if list_tag.name == 'ol' else 'unordered'
                    }
                ))

        # Process tables
        for table_idx, table in enumerate(soup.find_all('table')):
            rows = []
            for tr in table.find_all('tr'):
                cells = [td.get_text(strip=True) for td in tr.find_all(['td', 'th'])]
                if cells:
                    rows.append(cells)

            if rows:
                table_md = TableExtractor.table_to_markdown(rows)
                elements.append(DocumentElement(
                    type='table',
                    content=table_md,
                    metadata={
                        'table_index': table_idx,
                        'rows': len(rows),
                        'cols': len(rows[0]) if rows else 0,
                    }
                ))

        # Process code blocks
        for code in soup.find_all(['code', 'pre']):
            text = code.get_text()
            if text.strip():
                elements.append(DocumentElement(
                    type='code',
                    content=text,
                    metadata={'tag': code.name}
                ))

        combined_text = '\n\n'.join(full_text)

        return ParsedDocument(
            text=combined_text,
            elements=elements,
            metadata=metadata,
            structure={
                'total_elements': len(elements),
                'tables': sum(1 for e in elements if e.type == 'table'),
                'headings': sum(1 for e in elements if e.type == 'heading'),
                'code_blocks': sum(1 for e in elements if e.type == 'code'),
                'parser': 'html'
            },
            raw_content=content
        )


class TextParser:
    """Parse plain text documents."""

    def parse(self, file_path: Path) -> ParsedDocument:
        """
        Parse plain text document.

        Args:
            file_path: Path to text file

        Returns:
            Parsed document
        """
        # Read file
        text = EncodingDetector.read_text_file(file_path)

        # Clean text
        text = TextCleaner.clean_text(text)

        # Simple structure detection
        elements = self._parse_text_structure(text)

        # Extract metadata
        metadata = {
            'parser': 'text',
            'language': LanguageDetector.detect_language(text),
        }

        return ParsedDocument(
            text=text,
            elements=elements,
            metadata=metadata,
            structure={
                'total_elements': len(elements),
                'parser': 'text'
            },
            raw_content=text
        )

    def _parse_text_structure(self, text: str) -> List[DocumentElement]:
        """Parse plain text into structured elements."""
        elements = []
        paragraphs = re.split(r'\n\s*\n', text)

        for para in paragraphs:
            para = para.strip()
            if not para:
                continue

            # Check if it's a heading (short, ends without punctuation)
            if len(para) < 100 and not para.endswith(('.', ',', ';', ':', '!', '?')):
                if para[0].isupper():
                    elements.append(DocumentElement(
                        type='heading',
                        content=para,
                        level=1 if para.isupper() else 2
                    ))
                    continue

            # Check if it's a list
            list_pattern = r'^[\-\*\•\◦]\s+|^\d+[\.\)]\s+'
            lines = para.split('\n')
            if all(re.match(list_pattern, line.strip()) for line in lines if line.strip()):
                elements.append(DocumentElement(
                    type='list',
                    content=para,
                    metadata={'list_type': 'unordered'}
                ))
                continue

            # Regular paragraph
            elements.append(DocumentElement(
                type='paragraph',
                content=para
            ))

        return elements


class ImageParser:
    """Parse image documents using OCR."""

    def __init__(self, ocr_engine: Optional[Any] = None, enable_ocr: bool = True):
        """
        Initialize image parser.

        Args:
            ocr_engine: OCR engine instance
        """
        self.ocr_engine = ocr_engine
        self.enable_ocr = enable_ocr

    def parse(self, file_path: Path) -> ParsedDocument:
        """
        Parse image document.

        Args:
            file_path: Path to image file

        Returns:
            Parsed document with OCR results
        """
        if not self.enable_ocr or self.ocr_engine is None:
            raise RuntimeError("OCR engine is not configured; image documents cannot be processed")

        # Run OCR
        ocr_result = self.ocr_engine.process_image(file_path)

        # Parse structure from OCR result
        elements = []
        if ocr_result.text:
            # Simple paragraph detection
            paragraphs = re.split(r'\n\s*\n', ocr_result.text)
            for para in paragraphs:
                para = para.strip()
                if para:
                    elements.append(DocumentElement(
                        type='paragraph',
                        content=para,
                        metadata={'confidence': ocr_result.confidence}
                    ))

        # Extract metadata
        from .utils import ImagePreprocessor
        img_info = ImagePreprocessor.get_image_info(file_path)

        metadata = {
            'parser': 'ocr',
            'ocr_confidence': ocr_result.confidence,
            'processing_time': ocr_result.processing_time,
            'language': ocr_result.language,
            **img_info
        }

        return ParsedDocument(
            text=ocr_result.text,
            elements=elements,
            metadata=metadata,
            structure={
                'total_elements': len(elements),
                'ocr_boxes': len(ocr_result.boxes),
                'parser': 'ocr'
            },
            raw_content=ocr_result.text
        )


class DocumentParser:
    """
    Main document parser that routes to appropriate format-specific parser.

    Supports: PDF, DOCX, HTML, TXT, images (PNG, JPG, TIFF)
    """

    def __init__(self, ocr_engine: Optional[Any] = None, enable_ocr: bool = True):
        """
        Initialize document parser.

        Args:
            ocr_engine: OCR engine for scanned documents and images
        """
        self.enable_ocr = enable_ocr

        if enable_ocr:
            if ocr_engine is not None:
                self.ocr_engine = ocr_engine
            else:
                try:
                    self.ocr_engine = OCREngineFactory.create_engine()
                except Exception as err:
                    logger.warning(f"Failed to initialize OCR engine automatically: {err}")
                    self.ocr_engine = None
                    self.enable_ocr = False
        else:
            self.ocr_engine = None

        # Initialize format-specific parsers
        self.pdf_parser = PDFParser(ocr_engine=self.ocr_engine)
        self.docx_parser = DOCXParser()
        self.html_parser = HTMLParser()
        self.text_parser = TextParser()
        self.image_parser = ImageParser(
            ocr_engine=self.ocr_engine,
            enable_ocr=self.enable_ocr
        )

    def parse(self, file_path: Path) -> ParsedDocument:
        """
        Parse document of any supported format.

        Args:
            file_path: Path to document

        Returns:
            Parsed document with structure

        Raises:
            ValueError: If file format is not supported
        """
        if not file_path.exists():
            raise FileNotFoundError(f"File not found: {file_path}")

        # Detect file type
        file_type = FileTypeDetector.detect_file_type(file_path)

        logger.info(f"Parsing {file_type} document: {file_path.name}")

        # Route to appropriate parser
        if file_type == 'pdf':
            parsed = self.pdf_parser.parse(file_path)
        elif file_type in ['docx', 'doc']:
            parsed = self.docx_parser.parse(file_path)
        elif file_type in ['html', 'htm']:
            parsed = self.html_parser.parse(file_path)
        elif file_type in ['txt', 'md']:
            parsed = self.text_parser.parse(file_path)
        elif file_type in ['png', 'jpg', 'jpeg', 'tiff', 'bmp']:
            if not self.enable_ocr or self.ocr_engine is None:
                raise RuntimeError(
                    "OCR is disabled or unavailable; image documents cannot be parsed"
                )
            parsed = self.image_parser.parse(file_path)
        else:
            raise ValueError(f"Unsupported file format: {file_type}")

        # Add file metadata
        file_metadata = DocumentMetadataExtractor.extract_basic_metadata(file_path)
        parsed.metadata.update(file_metadata)

        # Add text statistics
        text_stats = DocumentMetadataExtractor.extract_text_statistics(parsed.text)
        parsed.metadata['statistics'] = text_stats

        logger.info(
            f"Parsed {file_path.name}: "
            f"{len(parsed.elements)} elements, "
            f"{text_stats['word_count']} words"
        )

        return parsed
